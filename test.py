import os
from docx import Document
from docx.shared import Pt

def compile_code_to_docx(target_folder=".", output_filename="compiled_code_for_ai.docx"):
    # Added .php to handle the LAMP stack
    code_extensions = {
        '.py', '.js', '.jsx', '.ts', '.tsx', '.html', '.css', 
        '.php', '.dart', '.c', '.cpp', '.h', '.java', '.json', '.yaml', '.yml', '.sh'
    }

    doc = Document()
    
    # Adjust default styling to make the Word doc visually dense
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(8)

    file_count = 0

    for root, dirs, files in os.walk(target_folder):
        # Ignore hidden directories like .git and build folders
        dirs[:] = [d for d in dirs if not d.startswith('.') and d not in ['node_modules', 'build', '__pycache__', 'vendor']]

        for file in files:
            ext = os.path.splitext(file)[1].lower()
            
            if ext in code_extensions:
                file_path = os.path.join(root, file)
                
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        lines = f.readlines()

                    # Token Minimization Strategy:
                    # 1. Remove entirely blank lines
                    # 2. Strip trailing whitespaces
                    # 3. Preserve leading spaces
                    processed_lines = [line.rstrip() for line in lines if line.strip()]

                    if not processed_lines:
                        continue  # Skip files that are empty or just whitespace

                    # Add a minimalistic header
                    rel_path = os.path.relpath(file_path, target_folder)
                    p_header = doc.add_paragraph(f"### {rel_path} ###")
                    p_header.paragraph_format.space_after = Pt(0)
                    p_header.paragraph_format.space_before = Pt(0)

                    # Join the processed lines and add to doc
                    code_text = "\n".join(processed_lines)
                    p_code = doc.add_paragraph(code_text)
                    p_code.paragraph_format.space_after = Pt(0)
                    p_code.paragraph_format.space_before = Pt(0)
                    
                    file_count += 1

                except Exception as e:
                    print(f"Skipping {file_path} due to error: {e}")

    doc.save(output_filename)
    print(f"Success! Compiled {file_count} files into {output_filename}.")

if __name__ == "__main__":
    current_directory = os.getcwd()
    print(f"Scanning directory: {current_directory}")
    compile_code_to_docx(target_folder=current_directory)